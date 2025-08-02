import os
import base64
import pdfplumber  # Alternative to PyMuPDF
from docx import Document
from PIL import Image
import io
import tempfile
from typing import Dict, List, Tuple, Optional

class FileProcessor:
    """Handles processing of PDFs, Word documents, and images"""
    
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self.process_pdf,
            '.docx': self.process_docx,
            '.doc': self.process_docx,
            '.jpg': self.process_image,
            '.jpeg': self.process_image,
            '.png': self.process_image,
            '.gif': self.process_image,
            '.bmp': self.process_image,
            '.webp': self.process_image
        }
    
    def get_file_extension(self, filename: str) -> str:
        """Get file extension from filename"""
        return os.path.splitext(filename.lower())[1]
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        ext = self.get_file_extension(filename)
        return ext in self.supported_extensions
    
    def process_file(self, file_data: bytes, filename: str) -> Dict:
        """Process any supported file type"""
        ext = self.get_file_extension(filename)
        
        if not self.is_supported_file(filename):
            return {
                "success": False,
                "error": f"Unsupported file type: {ext}",
                "content": None,
                "file_type": "unknown"
            }
        
        try:
            processor_func = self.supported_extensions[ext]
            result = processor_func(file_data, filename)
            result["file_type"] = ext[1:]  # Remove the dot
            return result
        except Exception as e:
            return {
                "success": False,
                "error": f"Error processing file: {str(e)}",
                "content": None,
                "file_type": ext[1:]
            }
    
    def process_pdf(self, file_data: bytes, filename: str) -> Dict:
        """Extract text from PDF using pdfplumber"""
        try:
            # Create a temporary file-like object
            pdf_stream = io.BytesIO(file_data)
            
            extracted_text = ""
            page_count = 0
            
            with pdfplumber.open(pdf_stream) as pdf:
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text from the page
                    text = page.extract_text()
                    if text:
                        extracted_text += f"\n--- Page {page_num + 1} ---\n{text}\n"
                    
                    # Extract tables if present
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables):
                            extracted_text += f"\n--- Table {table_num + 1} on Page {page_num + 1} ---\n"
                            for row in table:
                                if row and any(cell for cell in row if cell):
                                    row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                    extracted_text += row_text + "\n"
                            extracted_text += "\n"
            
            return {
                "success": True,
                "content": {
                    "text": extracted_text.strip(),
                    "page_count": page_count
                },
                "error": None
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF processing error: {str(e)}",
                "content": None
            }
    
    def process_docx(self, file_data: bytes, filename: str) -> Dict:
        """Extract text from Word document"""
        try:
            # Create a temporary file-like object
            doc_stream = io.BytesIO(file_data)
            doc = Document(doc_stream)
            
            extracted_text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        extracted_text += " | ".join(row_text) + "\n"
                extracted_text += "\n"  # Add space after table
            
            return {
                "success": True,
                "content": {
                    "text": extracted_text.strip(),
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables)
                },
                "error": None
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Word document processing error: {str(e)}",
                "content": None
            }
    
    def process_image(self, file_data: bytes, filename: str) -> Dict:
        """Process image and convert to base64"""
        try:
            # Open image with PIL to validate and get info
            image_stream = io.BytesIO(file_data)
            image = Image.open(image_stream)
            
            # Get image information
            image_info = {
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
                "width": image.width,
                "height": image.height
            }
            
            # Convert to base64
            image_b64 = base64.b64encode(file_data).decode('utf-8')
            
            # Determine MIME type
            mime_type_map = {
                'JPEG': 'image/jpeg',
                'PNG': 'image/png',
                'GIF': 'image/gif',
                'BMP': 'image/bmp',
                'WEBP': 'image/webp'
            }
            mime_type = mime_type_map.get(image.format, 'image/jpeg')
            
            return {
                "success": True,
                "content": {
                    "base64": image_b64,
                    "mime_type": mime_type,
                    "info": image_info,
                    "size_bytes": len(file_data)
                },
                "error": None
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Image processing error: {str(e)}",
                "content": None
            }
    
    def format_content_for_gpt(self, processed_files: List[Dict]) -> str:
        """Format processed files into a message for GPT"""
        if not processed_files:
            return ""
        
        formatted_content = []
        
        for file_info in processed_files:
            if not file_info.get("success"):
                continue
                
            filename = file_info.get("filename", "Unknown file")
            file_type = file_info.get("file_type", "unknown")
            content = file_info.get("content", {})
            
            if file_type in ['pdf', 'docx', 'doc']:
                # Text-based documents
                text = content.get("text", "")
                if text:
                    formatted_content.append(f"📄 **{filename}** ({file_type.upper()}):\n{text}\n")
                else:
                    formatted_content.append(f"📄 **{filename}** ({file_type.upper()}): No text content found\n")
                    
            elif file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp']:
                # Images
                info = content.get("info", {})
                size = info.get("size", (0, 0))
                formatted_content.append(f"🖼️ **{filename}** ({file_type.upper()}): Image {size[0]}x{size[1]} pixels\n")
        
        return "\n".join(formatted_content) 